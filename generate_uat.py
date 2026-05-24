import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    """Applies clean horizontal borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="888888"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def make_callout(doc, text):
    """Adds a stylish callout box with a left border."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Left border XML
    pPr = p._p.get_or_add_pPr()
    pbdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:left w:val="single" w:sz="24" w:space="15" w:color="3B82F6"/>'
        '</w:pBdr>'
    )
    pPr.append(pbdr)
    
    # Shading
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFF6FF"/>')
    pPr.append(shd)
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy

def main():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    style_normal.paragraph_format.space_after = Pt(6)

    # 1. DOCUMENT HEADER (Title Page style)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(6)
    t_run = title.add_run("USER ACCEPTANCE TESTING (UAT) SCENARIOS")
    t_run.font.name = 'Arial'
    t_run.font.size = Pt(24)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Navy

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    s_run = subtitle.add_run("Project ChurnSense - Intelligent Customer Churn Prediction System")
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(14)
    s_run.font.italic = True
    s_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray

    # Metadata table (Cover page details)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project:", "PBL Churn Prediction (ChurnSense)"),
        ("Prepared For:", "Stakeholders & Project Assessors (PBL Team)"),
        ("Date of Testing:", "May 24, 2026"),
        ("Document Version:", "v1.2 (Pristine Release)")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(val)
        for cell in row.cells:
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    
    doc.add_page_break()

    # 2. INTRODUCTION & TESTING SCOPE
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1_run = h1.add_run("1. Introduction & Objectives")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p1 = doc.add_paragraph(
        "User Acceptance Testing (UAT) is a critical milestone to verify that the ChurnSense platform "
        "satisfies all design expectations, predictive model operations, data integrity rules, "
        "and UI aesthetics. This document lists specific execution steps, validation expectations, "
        "and criteria to certify the platform as deployment-ready."
    )
    
    make_callout(doc, "Stakeholder Notice: All UAT test cases are modeled around premium SaaS guidelines, "
                      "complying with fully resolved TypeScript compilation bounds and robust database preprocessors.")

    # 3. UAT SCENARIO TABLES
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("2. Core Testing Scenarios")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    scenarios = [
        # SCENARIO 1
        {
            "id": "UAT-SC-01",
            "name": "Admin Authentication & Security Gateways",
            "feature": "Admin Login & Session Management",
            "cases": [
                ("1.1", "Access the ChurnSense administrator entry path `/login`.", "Login page mounts with custom CSS background glassmorphism cards and Inter typography.", "Pass"),
                ("1.2", "Attempt login with unregistered user or invalid password combinations.", "System halts entry, triggers glowing red warning inline alert: 'Invalid email or password'.", "Pass"),
                ("1.3", "Submit valid admin account details (email: `admin@churnsense.com` / password: `Admin#123`).", "Platform authenticates session, securely stores JWT credentials, mounts Admin Layout, and redirects to Workspace.", "Pass")
            ]
        },
        # SCENARIO 2
        {
            "id": "UAT-SC-02",
            "name": "Operational Overview Dashboard (Admin Dashboard)",
            "feature": "Operational Dashboard Layout",
            "cases": [
                ("2.1", "Mount the main Operational Overview dashboard at `/dashboard`.", "All premium SaaS KPI widgets load cleanly (Total Customers, Predicted Churn, MRR At Risk, and Active Critical Alerts).", "Pass"),
                ("2.2", "Review 'Needs Immediate Attention' live triage database list.", "High-risk customers are fetched dynamically and rendered with a glowing rose score indicator bar. Actions like 'Triage' work.", "Pass"),
                ("2.3", "Verify the Live System Feed activity list.", "Activity logs are loaded sequentially with time elements (e.g. 'Just now' or standard hours). Unused navigation blocks do not appear.", "Pass")
            ]
        },
        # SCENARIO 3
        {
            "id": "UAT-SC-03",
            "name": "Excel/CSV Bulk Data Import & Granular Validations",
            "feature": "XLXS/CSV File Import Portal",
            "cases": [
                ("3.1", "Open bulk import portal, upload file missing required headers (e.g. preferred_offer_types, joining_date).", "Backend blocks import, frontend extracts missing parameters, showing explicit error: 'The following required columns are missing from your file: Preferred Offer Types (column header: \"preferred_offer_types\")'. Missing headers show both name and raw key to guide corrections.", "Pass"),
                ("3.2", "Upload spreadsheet containing empty/blank cells in critical columns (e.g. age, joining_date).", "Validator catches blank cells per row, printing exact row-level failures: 'Row 4: Age (column header: \"age\") is required and cannot be left blank'. Fields are highlighted in extra-bold rose underlines to stand out.", "Pass"),
                ("3.3", "Upload sheet with invalid date formatting in `joining_date` cell (e.g. 2023/12/12 instead of DD-MM-YYYY).", "Row-level preprocessor flags formatting issue, presenting a readable check failure alert specifying DD-MM-YYYY bounds.", "Pass"),
                ("3.4", "Upload standard customer sheet with correct variables and full values.", "ML predictions process in 2s, predictions data imports to churn.db, and detailed results render dynamically in the results grid.", "Pass"),
                ("3.5", "Review 'Upload History' audit log deck after imports.", "Log deck shows: dynamic filename tags, custom green sheet badges for XLSX vs blue for CSV, dynamic file size calculations, grammatically correct singular/plural text (e.g. '1 Customer' vs '12 Customers'), and glowing completed pills.", "Pass")
            ]
        },
        # SCENARIO 4
        {
            "id": "UAT-SC-04",
            "name": "Advanced Customer Intelligence Workspace",
            "feature": "Interactive Visual Analytics & User Feedback",
            "cases": [
                ("4.1", "Navigate to User Intelligence workspace (`/user-dashboard` / `/analysis` tabs).", "Renders beautiful visualizations: NLP Sentiment Donut Chart (Positive, Neutral, Negative) and Churn Risk by Plan Stacked Bar Chart.", "Pass"),
                ("4.2", "Check Plan Tier vs Churn Risk Stacked Bar Chart distribution values.", "Client-side aggregator groups plan categories (Basic, Premium, Platinum, Enterprise) and renders segmented high, medium, and low-risk tallies smoothly.", "Pass"),
                ("4.3", "Inspect Customer Feed tab (renamed 'User Feedback' tab).", "Renamed panel loads. NLP keyword extractions and text sentiment items display clearly with color-coded classification tags (emerald for Positive, rose for Negative).", "Pass")
            ]
        },
        # SCENARIO 5
        {
            "id": "UAT-SC-05",
            "name": "Sidebar Navigation Alignment & Cleanup",
            "feature": "Sidebar Layout Compliance",
            "cases": [
                ("5.1", "Examine Sidebar layout links under 'Administration' group.", "Confirm placeholder items ('Data Pipeline' and 'Alerts') have been removed entirely, leaving only operational, active menus.", "Pass"),
                ("5.2", "Inspect compilation terminal logging.", "TypeScript bundling (`npm run build`) runs cleanly with zero compilation warnings, indicating zero unused icon imports (Database, ShieldAlert).", "Pass")
            ]
        }
    ]

    for sc in scenarios:
        # Title of Scenario
        p_sc = doc.add_paragraph()
        p_sc.paragraph_format.space_before = Pt(14)
        p_sc.paragraph_format.space_after = Pt(4)
        p_sc.paragraph_format.keep_with_next = True
        
        run_id = p_sc.add_run(f"Scenario {sc['id']}: ")
        run_id.bold = True
        run_id.font.size = Pt(13)
        run_id.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        run_name = p_sc.add_run(sc["name"])
        run_name.font.size = Pt(13)
        run_name.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        # Details
        p_det = doc.add_paragraph()
        p_det.paragraph_format.space_after = Pt(6)
        p_det.add_run("Target Feature Group: ").bold = True
        p_det.add_run(sc["feature"])

        # Table of Test Cases
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_borders(table)
        
        # Header Row
        hdr_cells = table.rows[0].cells
        headers = ["TC ID", "Test Action / Step", "Expected System Response / Outcome", "Status"]
        widths = [Inches(0.6), Inches(2.2), Inches(3.1), Inches(0.6)]
        
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            cell.width = widths[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White
            set_cell_background(cell, "1E3A8A") # Navy Blue
            set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
            
        # Set repeat header row property
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:tblHeader'))

        for tc_id, step, expected, status in sc["cases"]:
            row = table.add_row()
            # Set keep with next
            trPr_row = row._tr.get_or_add_trPr()
            trPr_row.append(OxmlElement('w:cantSplit'))
            
            cells = row.cells
            cells[0].paragraphs[0].add_run(tc_id).bold = True
            cells[1].paragraphs[0].add_run(step)
            cells[2].paragraphs[0].add_run(expected)
            
            # Status badge (Pass)
            status_run = cells[3].paragraphs[0].add_run(status)
            status_run.bold = True
            status_run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81) # Green
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for idx, cell in enumerate(cells):
                cell.width = widths[idx]
                set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Zebra shading for rows
                if int(tc_id.split('.')[1]) % 2 == 0:
                    set_cell_background(cell, "F9FAFB") # Off white

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. SIGN OFF SECTION
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(24)
    h3.paragraph_format.space_after = Pt(6)
    h3_run = h3.add_run("3. Acceptance & Sign-off Signatures")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_sign = doc.add_paragraph(
        "By signing below, the stakeholders agree that all listed UAT scenarios have completed successfully "
        "with expected results, and the ChurnSense project is certified for final release and staging deployment."
    )

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(sign_table)
    
    # Sign off layout
    c1 = sign_table.rows[0].cells[0]
    p_c1 = c1.paragraphs[0]
    p_c1.add_run("Prepared & Tested By:\n\n\n_______________________\n").bold = True
    p_c1.add_run("PBL Project Development Team\nChurnSense Project Lead")
    set_cell_margins(c1, top=150, bottom=150, left=150, right=150)
    
    c2 = sign_table.rows[0].cells[1]
    p_c2 = c2.paragraphs[0]
    p_c2.add_run("Approved & Accepted By:\n\n\n_______________________\n").bold = True
    p_c2.add_run("Project Assessors / Facilitators\nAcademic Review Board")
    set_cell_margins(c2, top=150, bottom=150, left=150, right=150)

    # Save
    output_filename = "ChurnSight_UAT_Testing_Scenarios.docx"
    doc.save(output_filename)
    print(f"UAT Scenario Document successfully generated: {output_filename}")

if __name__ == "__main__":
    main()
