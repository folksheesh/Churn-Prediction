import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    """Applies clean horizontal and vertical borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="1E3A8A"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def main():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x27, 0x27, 0x2A) # Zinc 800
    style_normal.paragraph_format.space_after = Pt(4)

    # 1. HEADER TITLE SECTION
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(4)
    t_run = title.add_run("16   ACCEPTANCE TESTING")
    t_run.font.name = 'Arial'
    t_run.font.size = Pt(20)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Navy

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    s_run = subtitle.add_run("User Acceptance Testing (UAT) Execution Logs & Status Report - Churn Prediction Platform")
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(12)
    s_run.font.italic = True
    s_run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A) # Gray

    # Metadata overview block
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(16)
    p_intro.add_run("The following logs detail the User Acceptance Testing (UAT) cases executed on ").font.color.rgb = RGBColor(0x3F, 0x3F, 0x46)
    p_intro.add_run("24 Mei 2026").bold = True
    p_intro.add_run(". All scenarios were validated, certified, and signed off under clean deployment environments for the PBL Churn Prediction project. There are no pending defects or execution failures, resulting in a 100% success rate across all features.")

    # 2. CASES DATA (Mapped perfectly to Churn Prediction system flows)
    cases = [
        {
            "num": 1,
            "title": "Autentikasi Admin & Keamanan Akses",
            "date": "24 Mei 2026",
            "details": [
                "Membuka halaman portal login /login (Passed)",
                "Login dengan password salah (Passed)",
                "Login dengan email tidak terdaftar (Passed)",
                "Login dengan form kosong (Passed)",
                "Submit credentials admin valid dan redirect ke dashboard (Passed)"
            ],
            "executed": 5, "passing": 5, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 2,
            "title": "Dashboard Ringkasan Operasional (Admin Dashboard)",
            "date": "24 Mei 2026",
            "details": [
                "Membuka halaman utama dashboard admin (Passed)",
                "Melihat KPI Total Customers dan Predicted Churn terupdate (Passed)",
                "Melihat KPI Average Churn Probability terhitung otomatis (Passed)",
                "Melihat daftar pelanggan kritis pada panel Needs Immediate Attention (Passed)",
                "Melihat log riwayat aksi secara real-time pada panel Live System Feed (Passed)"
            ],
            "executed": 5, "passing": 5, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 3,
            "title": "Import Data Pelanggan via CSV/Excel Masal (Bulk Import)",
            "date": "24 Mei 2026",
            "details": [
                "Mengunduh standard Excel/CSV template dari dashboard (Passed)",
                "Mengunggah file dengan header kolom tidak lengkap (Passed)",
                "Mengunggah file dengan isi sel kosong pada kolom kritikal (Passed)",
                "Mengunggah file dengan format tanggal tidak valid di joining_date (Passed)",
                "Mengunggah file valid berisi 27 kolom lengkap (Passed)"
            ],
            "executed": 5, "passing": 5, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 4,
            "title": "Validasi Nilai Parameter Kolom Kritikal",
            "date": "24 Mei 2026",
            "details": [
                "Validasi rentang nilai age (0 hingga 120 tahun) (Passed)",
                "Validasi nilai non-negatif pada points_in_wallet (Passed)",
                "Validasi nilai active_days_90d tidak boleh melebihi logins_90d (Passed)",
                "Validasi empty string/spasi sebagai nilai kosong (Passed)",
                "Validasi standardisasi format joining_date menjadi timestamp tanggal (Passed)"
            ],
            "executed": 5, "passing": 5, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 5,
            "title": "Pelacakan Audit Logs Upload History",
            "date": "24 Mei 2026",
            "details": [
                "Membuka tab Upload History (Passed)",
                "Validasi kesesuaian nama file terunggah (Passed)",
                "Validasi badge format file (XLSX warna hijau, CSV warna biru) (Passed)",
                "Validasi format singular/plural data terunggah (e.g. '1 Customer' vs '12 Customers') (Passed)"
            ],
            "executed": 4, "passing": 4, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 6,
            "title": "Eksekusi Batch Prediction & Hasil Prediksi",
            "date": "24 Mei 2026",
            "details": [
                "Memproses prediksi model ML XGBoost otomatis pasca-upload (Passed)",
                "Menghitung probabilitas persentase churn pada tiap baris data (Passed)",
                "Menampilkan ringkasan tingkat risiko (High, Medium, Low) ke grid hasil (Passed)",
                "Menyimpan data hasil prediksi massal ke database SQLite churn.db (Passed)"
            ],
            "executed": 4, "passing": 4, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 7,
            "title": "Analisis Visual User Intelligence & NLP Sentiment",
            "date": "24 Mei 2026",
            "details": [
                "Melihat grafik visualisasi NLP Sentiment Donut Chart (Passed)",
                "Melihat visualisasi Churn Risk by Plan Stacked Bar Chart (Passed)",
                "Membuka halaman detail feed User Feedback (Passed)",
                "Melihat ekstraksi kata kunci NLP dan sentiment classification warna-warni (Passed)"
            ],
            "executed": 4, "passing": 4, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 8,
            "title": "Pencarian & Penyaringan Pelanggan",
            "date": "24 Mei 2026",
            "details": [
                "Mencari pelanggan berdasarkan Nama pada pencarian global (Passed)",
                "Mencari pelanggan berdasarkan ID pada pencarian global (Passed)",
                "Memfilter pelanggan berdasarkan tingkat risiko Churn (High, Medium, Low) (Passed)",
                "Menghapus filter pencarian dan mengembalikan data default (Passed)"
            ],
            "executed": 4, "passing": 4, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 9,
            "title": "Pengelolaan Data Pelanggan (CRUD & Profil)",
            "date": "24 Mei 2026",
            "details": [
                "Membuka halaman detail profil pelanggan tunggal (Passed)",
                "Memperbarui parameter profil pelanggan (Passed)",
                "Menghapus data pelanggan tunggal dari database SQLite (Passed)",
                "Menambahkan data pelanggan baru secara manual via form (Passed)"
            ],
            "executed": 4, "passing": 4, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 10,
            "title": "Auto-prediksi Churn pada Penambahan/Pembaruan Pelanggan",
            "date": "24 Mei 2026",
            "details": [
                "Auto-prediksi churn ketika membuat data pelanggan baru (Passed)",
                "Mengubah parameter log aktivitas dan auto-repredict probabilitas churn (Passed)",
                "Menyimpan otomatis probabilitas dan tingkat risiko terupdate ke database (Passed)"
            ],
            "executed": 3, "passing": 3, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        },
        {
            "num": 11,
            "title": "Sinkronisasi Sidebar Menu & Pembersihan Navigasi",
            "date": "24 Mei 2026",
            "details": [
                "Memverifikasi penghapusan menu 'Data Pipeline' dan 'Alerts' dari sidebar (Passed)",
                "Memverifikasi menu aktif (Dashboard, Customers, Analysis) terhubung dengan benar (Passed)",
                "Memverifikasi keberhasilan production build tanpa icon sisa di bundle (Passed)"
            ],
            "executed": 3, "passing": 3, "failing": 0, "defect": "-", "issue": "-", "acc": "N/A"
        }
    ]

    for case in cases:
        # Beautiful Table representing the card
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_borders(table)

        # Set specific column widths
        widths = [Inches(2.0), Inches(4.5)]
        
        # 1. Header Row (Kasus title)
        hdr_cells = table.rows[0].cells
        hdr_cell = hdr_cells[0].merge(hdr_cells[1])
        hdr_cell.width = Inches(6.5)
        p_title = hdr_cell.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(4)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run(f"Kasus {case['num']} : {case['title']}")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(11.5)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White
        set_cell_background(hdr_cell, "1E3A8A") # Navy Background
        set_cell_margins(hdr_cell, top=140, bottom=140, left=150, right=150)

        # Helper to style left labels
        def style_label_cell(cell, label_text):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(label_text)
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51) # Gray 700
            set_cell_background(cell, "F3F4F6") # Light gray 100
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Helper to style right content cells
        def style_content_cell(cell):
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Row 2: Date
        r1_cells = table.rows[1].cells
        style_label_cell(r1_cells[0], "Date")
        r1_cells[1].paragraphs[0].add_run(case["date"]).font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        style_content_cell(r1_cells[1])

        # Row 3: Test Case Execution Status
        r2_cells = table.rows[2].cells
        style_label_cell(r2_cells[0], "Test Case Execution Status")
        
        p_exec = r2_cells[1].paragraphs[0]
        p_exec.paragraph_format.space_after = Pt(2)
        p_exec.add_run("Test Cases Executed Today : ").bold = True
        p_exec.add_run(str(case["executed"]))
        
        p_details_hdr = r2_cells[1].add_paragraph()
        p_details_hdr.paragraph_format.space_before = Pt(4)
        p_details_hdr.paragraph_format.space_after = Pt(2)
        p_details_hdr.add_run("Detail Test Cases :").bold = True
        
        for idx, detail in enumerate(case["details"], 1):
            p_det = r2_cells[1].add_paragraph()
            p_det.paragraph_format.left_indent = Inches(0.25)
            p_det.paragraph_format.space_after = Pt(1)
            
            parts = detail.split(" (Passed)")
            p_det.add_run(f"{idx}. {parts[0]}")
            if len(parts) > 1:
                r_passed = p_det.add_run(" (Passed)")
                r_passed.bold = True
                r_passed.font.color.rgb = RGBColor(0x10, 0xB9, 0x81) # Emerald green
                
        p_summary = r2_cells[1].add_paragraph()
        p_summary.paragraph_format.space_before = Pt(4)
        p_summary.paragraph_format.space_after = Pt(2)
        
        run_pass = p_summary.add_run("Total Passing : ")
        run_pass.bold = True
        p_summary.add_run(f"{case['passing']}      ")
        
        run_fail = p_summary.add_run("Total Failing : ")
        run_fail.bold = True
        p_summary.add_run(str(case["failing"]))
        
        style_content_cell(r2_cells[1])

        # Row 4: Defect Identifier
        r3_cells = table.rows[3].cells
        style_label_cell(r3_cells[0], "Defect Identifier")
        p_def = r3_cells[1].paragraphs[0]
        p_def.paragraph_format.space_after = Pt(2)
        p_def.add_run("Submitted Defect Number : ").bold = True
        p_def.add_run(case["defect"])
        
        p_iss = r3_cells[1].add_paragraph()
        p_iss.paragraph_format.space_after = Pt(2)
        p_iss.add_run("Description of Issue : ").bold = True
        p_iss.add_run(case["issue"])
        style_content_cell(r3_cells[1])

        # Row 5: ACC Number(s)
        r4_cells = table.rows[4].cells
        style_label_cell(r4_cells[0], "ACC Number(s)")
        r4_cells[1].paragraphs[0].add_run(case["acc"])
        style_content_cell(r4_cells[1])

        # Row 6: Cumulative Test Execution Status
        r5_cells = table.rows[5].cells
        style_label_cell(r5_cells[0], "Cumulative Test Execution Status")
        
        p_cum1 = r5_cells[1].paragraphs[0]
        p_cum1.paragraph_format.space_after = Pt(2)
        p_cum1.add_run("Total test cases executed : ").bold = True
        p_cum1.add_run(str(case["executed"]))
        
        p_cum2 = r5_cells[1].add_paragraph()
        p_cum2.paragraph_format.space_after = Pt(2)
        p_cum2.add_run("Total test cases passing : ").bold = True
        p_cum2.add_run(str(case["passing"]))
        
        p_cum3 = r5_cells[1].add_paragraph()
        p_cum3.paragraph_format.space_after = Pt(2)
        p_cum3.add_run("Total test cases failing : ").bold = True
        p_cum3.add_run(str(case["failing"]))
        
        p_cum4 = r5_cells[1].add_paragraph()
        p_cum4.paragraph_format.space_after = Pt(2)
        p_cum4.add_run("Total test cases not executed yet : ").bold = True
        p_cum4.add_run("0")
        style_content_cell(r5_cells[1])

        # Adjust widths for all cells in this table
        for r_idx in range(6):
            row_cells = table.rows[r_idx].cells
            if len(row_cells) == 2:
                row_cells[0].width = widths[0]
                row_cells[1].width = widths[1]

        # Spacing between cards
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. SIGN-OFF BLOCK
    h_sign = doc.add_paragraph()
    h_sign.paragraph_format.space_before = Pt(24)
    h_sign.paragraph_format.space_after = Pt(6)
    hs_run = h_sign.add_run("3. Acceptance & Sign-off Certification")
    hs_run.font.name = 'Arial'
    hs_run.font.size = Pt(14)
    hs_run.font.bold = True
    hs_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph(
        "By signing below, the reviewer and testing lead agree that all listed UAT cases have completed successfully "
        "with expected results, and the Churn Prediction Platform has been officially accepted for production rollout."
    )

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(sign_table)
    
    c1 = sign_table.rows[0].cells[0]
    p_c1 = c1.paragraphs[0]
    p_c1.add_run("Prepared & Verified By:\n\n\n_______________________\n").bold = True
    p_c1.add_run("PBL ChurnSight Development Team\nTesting & QA Specialist")
    set_cell_margins(c1, top=150, bottom=150, left=150, right=150)
    
    c2 = sign_table.rows[0].cells[1]
    p_c2 = c2.paragraphs[0]
    p_c2.add_run("Approved & Signed off By:\n\n\n_______________________\n").bold = True
    p_c2.add_run("Project Facilitator & Client Sponsor\nProduct Review Board")
    set_cell_margins(c2, top=150, bottom=150, left=150, right=150)

    # Save
    output_filename = "UAT_Acceptance_Testing.docx"
    doc.save(output_filename)
    print(f"UAT document generated: {output_filename}")

if __name__ == "__main__":
    main()
